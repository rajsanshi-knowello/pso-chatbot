import os
import tempfile
from typing import Literal

import httpx
from fastapi import HTTPException
from pydantic import BaseModel

_MAX_BYTES = 10 * 1024 * 1024  # 10 MB
_DOWNLOAD_TIMEOUT = 30.0
_HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5", "Heading 6"}


class ParsedDocument(BaseModel):
    text: str
    word_count: int
    paragraph_count: int
    heading_count: int
    file_type: Literal["docx", "pdf"]
    has_tables: bool


def detect_file_type(name: str, content_type: str) -> Literal["docx", "pdf"]:
    lower = name.lower().split("?")[0]
    if lower.endswith(".docx"):
        return "docx"
    if lower.endswith(".pdf"):
        return "pdf"
    if "pdf" in content_type:
        return "pdf"
    if "wordprocessingml" in content_type or "docx" in content_type:
        return "docx"
    raise HTTPException(status_code=415, detail="Unsupported file type. Only .docx and .pdf are accepted.")


def _parse_docx(path: str) -> ParsedDocument:
    from docx import Document  # lazy import

    try:
        doc = Document(path)
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Could not parse DOCX file: {exc}") from exc

    text_parts: list[str] = []
    heading_count = 0
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
            if para.style.name in _HEADING_STYLES:
                heading_count += 1

    text = "\n\n".join(text_parts)
    return ParsedDocument(
        text=text,
        word_count=len(text.split()),
        paragraph_count=len(text_parts),
        heading_count=heading_count,
        file_type="docx",
        has_tables=len(doc.tables) > 0,
    )


def _parse_pdf(path: str) -> ParsedDocument:
    from pypdf import PdfReader  # lazy import

    try:
        reader = PdfReader(path)
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Could not parse PDF file: {exc}") from exc

    pages: list[str] = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(page_text)

    text = "\n\n".join(pages)
    paragraph_count = len([line for line in text.splitlines() if line.strip()])
    return ParsedDocument(
        text=text,
        word_count=len(text.split()),
        paragraph_count=paragraph_count,
        heading_count=0,
        file_type="pdf",
        has_tables=False,
    )


async def parse_document(url: str) -> ParsedDocument:
    try:
        async with httpx.AsyncClient(timeout=_DOWNLOAD_TIMEOUT, follow_redirects=True) as client:
            response = await client.get(url)
    except httpx.TimeoutException as exc:
        raise HTTPException(status_code=502, detail="Timed out downloading document.") from exc
    except httpx.RequestError as exc:
        raise HTTPException(status_code=502, detail=f"Could not reach document URL: {exc}") from exc

    if response.status_code != 200:
        raise HTTPException(
            status_code=502,
            detail=f"Document URL returned HTTP {response.status_code}.",
        )

    content_length = int(response.headers.get("content-length", 0))
    if content_length > _MAX_BYTES:
        raise HTTPException(status_code=413, detail="File exceeds the 10 MB limit.")

    content = response.content
    if len(content) > _MAX_BYTES:
        raise HTTPException(status_code=413, detail="File exceeds the 10 MB limit.")

    file_type = detect_file_type(url, response.headers.get("content-type", ""))

    with tempfile.NamedTemporaryFile(suffix=f".{file_type}", delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        if file_type == "docx":
            return _parse_docx(tmp_path)
        return _parse_pdf(tmp_path)
    finally:
        os.unlink(tmp_path)


def parse_document_bytes(content: bytes, filename: str, content_type: str) -> ParsedDocument:
    if len(content) > _MAX_BYTES:
        raise HTTPException(status_code=413, detail="File exceeds the 10 MB limit.")

    file_type = detect_file_type(filename, content_type)

    with tempfile.NamedTemporaryFile(suffix=f".{file_type}", delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        if file_type == "docx":
            return _parse_docx(tmp_path)
        return _parse_pdf(tmp_path)
    finally:
        os.unlink(tmp_path)
